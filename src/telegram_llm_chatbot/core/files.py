import base64
import io
import os
from typing import Set

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from PIL import Image
from PyPDF2 import PdfReader

from telegram_llm_chatbot.core.exceptions import (
    FileTooLargeException,
    PDFFileReadingException,
    TextFileDecodingException,
    UnexpectedFileReadingException,
    UnsupportedFileTypeException,
    WordFileReadingException,
)


class TextFileParser:
    """Class to parse and extract content from uploaded files."""

    def __init__(self, max_file_size_mb: int, allowed_file_types: Set[str]):
        """Initialize the TextFileParser."""
        self.max_file_size_mb = max_file_size_mb
        self.handlers = {
            "txt": self.extract_txt_content,
            "doc": self.extract_word_content,
            "docx": self.extract_word_content,
            "pdf": self.extract_pdf_content,
        }
        self.allowed_file_types = allowed_file_types

    def extract_txt_content(self, file_path: str) -> str:
        """Extract content from a text file.

        Returns:
            The content of the text file.

        Raises:
            TextFileDecodingException: Error decoding the text file.
            UnexpectedFileReadingException: Unexpected error while reading the text file.
        """
        try:
            with open(file_path, encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError as e:
            raise TextFileDecodingException() from e
        except Exception as e:
            raise UnexpectedFileReadingException() from e

    def extract_word_content(self, file_path: str) -> str:
        """Extract content from a Word document.

        Extract the text content from paragraphs and
        tables in the Word document while preserving the order of elements.

        Returns:
            The content of the Word document.
        Raises:
            WordFileReadingException: Error reading the Word document.
        """
        try:
            doc = docx.Document(file_path)
            content = []

            for element in doc.element.body:
                # Check if the element is a paragraph (CT_P)
                if isinstance(element, CT_P):
                    for para in doc.paragraphs:
                        if para._element == element:
                            paragraph_text = para.text.strip()
                            if paragraph_text:
                                content.append(paragraph_text)
                            break

                # Check if the element is a table (CT_Tbl)
                elif isinstance(element, CT_Tbl):
                    for table in doc.tables:
                        if table._element == element:
                            for row in table.rows:
                                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                                content.append(row_text)
                            content.append("")
                            break

            return "\n".join(content)
        except Exception as e:
            raise WordFileReadingException() from e

    def extract_pdf_content(self, file_path: str) -> str:
        """Extract content from a PDF file.

        Returns:
            The content of the PDF file.

        Raises:
            PDFFileReadingException: Error reading the PDF file.
        """
        try:
            pdf = PdfReader(file_path)
            content = ""
            for page in pdf.pages:
                content += page.extract_text()
            return content
        except Exception as e:
            raise PDFFileReadingException() from e

    def extract_content(self, file_path: str) -> str:
        """Extract content from an uploaded file based on its type.

        Returns:
            The content of the file.

        Raises:
            FileTooLargeException: The file is too large.
            UnsupportedFileTypeException: The file type is not supported.
            TextFileDecodingException: Error decoding the text file.
            WordFileReadingException: Error reading the Word document.
            UnexpectedFileReadingException: Unexpected error while reading the file.
        """
        if os.path.getsize(file_path) > self.max_file_size_mb * 1024 * 1024:
            raise FileTooLargeException()

        try:
            file_extension = file_path.rsplit(".", 1)[1].lower()
        except IndexError:
            raise UnsupportedFileTypeException("The file has no extension")  # noqa: WPS326

        handler = self.handlers.get(file_extension)
        if handler:
            try:
                return handler(file_path)
            except UnsupportedFileTypeException as e:
                raise e
            except TextFileDecodingException as e:
                raise e
            except WordFileReadingException as e:
                raise e
            except Exception as e:
                raise UnexpectedFileReadingException(f"Unexpected error: {str(e)}") from e
        else:
            raise UnsupportedFileTypeException(file_extension)


def image_to_base64(image: Image) -> str:
    """
    Converts a PIL Image to a base64 string.

    Args:
        image (Image): The image to convert.

    Returns:
        str: Base64 encoded string of the image.
    """
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")  # or any other format
    return base64.b64encode(buffered.getvalue()).decode()
